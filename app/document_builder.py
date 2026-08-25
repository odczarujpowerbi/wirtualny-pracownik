"""
Silnik budowy dokumentów narracyjnych (PDF/DOCX/MD) — rozszerzenie
report_builder.py (który robi tylko tabele markdown/CSV/XLSX) na dokumenty ze
strukturą: tytuł, sekcje z nagłówkiem, tekstem akapitowym i/lub tabelą.

Powód osobnego pliku, nie dopisania do report_builder.py: report_builder.py
jest CELOWO bez zależności zewnętrznych do markdown/CSV; PDF i DOCX wymagają
pakietów (`reportlab`, `python-docx`, patrz requirements.txt), więc trzymamy
je w module, który dało się nie importować, gdy pakietów nie ma.

Decyzja właściciela (22.08.2026): NIGDY .qmd (Quarto) — nikt tu nie używa
Quarto, plik nie otwiera się wygodnie od ręki. Zestawienia liczbowe (cokolwiek
do sumowania/filtrowania) ZAWSZE jako .xlsx (report_builder.write_xlsx_report),
nie jako tabela w dokumencie tekstowym. `.md` tylko jako lżejszy wariant, gdy
PDF/DOCX jest przesadą (krótka notatka, brak tabeli). Domyślny wybór dla
treści narracyjnej: PDF albo DOCX.

To jest silnik (przyjmuje gotową treść, formatuje wynik) — nie źródło
prawdy o KONKRETNYM dokumencie. Podłączenie realnych danych (CRM, MailerLite,
Meta Ads) to osobny krok, wykonywany przez skrypt, który wywołuje te funkcje.

Import `docx` i `reportlab` jest LENIWY (wewnątrz funkcji) — ten sam wzorzec
co `microsoft_graph_mail_client.py` — dzięki temu samo zaimportowanie tego
modułu nie wymaga obecności pakietów.

Domyślny katalog wyjściowy: runs/documents/ (gitignored, patrz .gitignore
"runs/") — ten sam wzorzec co screenshot_capture.py (runs/screenshots/).

Struktura `sections` (współdzielona przez build_docx, build_pdf i build_md) —
lista słowników, każdy jedno z:
  {"heading": "Nagłówek", "text": "Akapit...\n\nDrugi akapit..."}
  {"heading": "Nagłówek", "table": {"rows": [...], "columns": [...] | None}}
"""

import re
from pathlib import Path

import report_builder as _rb

DEFAULT_OUTPUT_DIR = Path(__file__).parent / "runs" / "documents"

# reportlab bez rejestracji własnego fontu używa Helvetica/Times (14 fontów
# bazowych PDF) — te NIE obejmują polskich znaków diakrytycznych (ą/ć/ę/ł/ń/
# ó/ś/ź/ż), więc wychodzą jako puste kwadraty (żywy incydent 25.08.2026:
# "Sprzeda■ Konferencji" w realnym raporcie MailerLite). Kandydaci fontów z
# pełnym pokryciem Latin Extended-A, w kolejności prawdopodobieństwa
# znalezienia na maszynie (Windows -> Linux z fontami DejaVu).
_PDF_FONT_KANDYDACI = (
    (Path(r"C:\Windows\Fonts\arial.ttf"), Path(r"C:\Windows\Fonts\arialbd.ttf")),
    (Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
     Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf")),
)
_PDF_FONT_NAME = "DokumentPL"

# Emoji w tytułach kampanii (np. z MailerLite) nie mają odpowiednika w
# zwykłym foncie tekstowym (kolorowe glify wymagają osobnego fontu bitmapowego,
# którego reportlab nie renderuje) — wychodziłyby jako te same puste kwadraty
# co brak polskich znaków. W PDF-ie po prostu je pomijamy; w .md/.docx
# zostają, bo tam renderuje je czcionka systemowa czytelnika.
_EMOJI_PATTERN = re.compile(
    "["
    "\U0001F300-\U0001FAFF"
    "\U00002600-\U000027BF"
    "\U0001F1E6-\U0001F1FF"
    "\U00002190-\U000021FF"
    "\U00002B00-\U00002BFF"
    "\uFE0F"
    "]+",
    flags=re.UNICODE,
)


def _bez_emoji(text):
    return _EMOJI_PATTERN.sub("", str(text)).strip()


def _zarejestruj_font_pdf():
    """Rejestruje pierwszy znaleziony font z kandydatów jako `_PDF_FONT_NAME`
    (+ wariant Bold). Zwraca nazwę fontu do użycia w stylach, albo None, gdy
    żadnego kandydata nie znaleziono — wtedy PDF i tak powstaje, tylko ze
    zdegradowanymi polskimi znakami (fail-soft, jak reszta modułu)."""
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    if _PDF_FONT_NAME in pdfmetrics.getRegisteredFontNames():
        return _PDF_FONT_NAME
    for regular, bold in _PDF_FONT_KANDYDACI:
        if not regular.exists():
            continue
        pdfmetrics.registerFont(TTFont(_PDF_FONT_NAME, str(regular)))
        pdfmetrics.registerFont(TTFont(_PDF_FONT_NAME + "-Bold", str(bold if bold.exists() else regular)))
        return _PDF_FONT_NAME
    return None


def _ensure_parent(output_path):
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    return output_path


def build_md(title, sections, output_path):
    """Plik Markdown zwykły (.md) — biblioteka standardowa, żadnej zależności.
    Dla lekkich notatek bez liczb do sumowania; zestawienia liczbowe -> .xlsx
    (report_builder.write_xlsx_report), nie tabela w tym pliku."""
    lines = [f"# {title}", ""]
    for section in sections:
        heading = section.get("heading")
        if heading:
            lines.append(f"## {heading}")
            lines.append("")
        text = section.get("text")
        if text:
            lines.append(text)
            lines.append("")
        table_spec = section.get("table")
        if table_spec:
            rows = table_spec.get("rows", [])
            columns = table_spec.get("columns")
            lines.append(_rb.build_markdown_table(rows, columns))
            lines.append("")

    output_path = _ensure_parent(output_path)
    output_path.write_text("\n".join(lines), encoding="utf-8")
    return output_path


def build_docx(title, sections, output_path):
    """Dokument Word (.docx) z tytułem i sekcjami (nagłówek + tekst i/lub tabela)."""
    import docx

    document = docx.Document()
    document.add_heading(title, level=0)

    for section in sections:
        heading = section.get("heading")
        if heading:
            document.add_heading(heading, level=1)

        text = section.get("text")
        if text:
            for paragraph in text.split("\n\n"):
                document.add_paragraph(paragraph)

        table_spec = section.get("table")
        if table_spec:
            rows = table_spec.get("rows", [])
            columns = table_spec.get("columns") or (list(rows[0].keys()) if rows else [])
            if columns:
                table = document.add_table(rows=1, cols=len(columns))
                table.style = "Light Grid Accent 1"
                header_cells = table.rows[0].cells
                for i, column in enumerate(columns):
                    header_cells[i].text = str(column)
                for row in rows:
                    cells = table.add_row().cells
                    for i, column in enumerate(columns):
                        cells[i].text = str(row.get(column, ""))

    output_path = _ensure_parent(output_path)
    document.save(str(output_path))
    return output_path


def build_pdf(title, sections, output_path):
    """Dokument PDF (.pdf) z tytułem i sekcjami (nagłówek + tekst i/lub tabela).
    Layout Platypus (reportlab) — proste flowables, bez projektowania szablonu."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    output_path = _ensure_parent(output_path)
    styles = getSampleStyleSheet()
    font = _zarejestruj_font_pdf()
    if font:
        styles["Title"].fontName = font + "-Bold"
        styles["Heading2"].fontName = font + "-Bold"
        styles["BodyText"].fontName = font
    doc = SimpleDocTemplate(str(output_path), pagesize=A4)
    flowables = [Paragraph(_bez_emoji(title), styles["Title"]), Spacer(1, 12)]

    for section in sections:
        heading = section.get("heading")
        if heading:
            flowables.append(Paragraph(_bez_emoji(heading), styles["Heading2"]))
            flowables.append(Spacer(1, 6))

        text = section.get("text")
        if text:
            for paragraph in text.split("\n\n"):
                flowables.append(Paragraph(_bez_emoji(paragraph).replace("\n", "<br/>"), styles["BodyText"]))
                flowables.append(Spacer(1, 6))

        table_spec = section.get("table")
        if table_spec:
            rows = table_spec.get("rows", [])
            columns = table_spec.get("columns") or (list(rows[0].keys()) if rows else [])
            if columns:
                data = [[_bez_emoji(c) for c in columns]] + \
                       [[_bez_emoji(row.get(c, "")) for c in columns] for row in rows]
                table = Table(data, hAlign="LEFT")
                table.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2f3b52")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("FONTNAME", (0, 0), (-1, -1), font or "Helvetica"),
                ]))
                flowables.append(table)
                flowables.append(Spacer(1, 12))

    doc.build(flowables)
    return output_path


if __name__ == "__main__":
    demo_sections = [
        {"heading": "Kontekst", "text": "Przykladowy dokument demonstracyjny document_builder.py."},
        {"heading": "Zestawienie", "table": {"rows": [
            {"Kanal": "Newsletter", "Wyslane": 1200, "Otwarcia": 480},
            {"Kanal": "Meta Ads", "Wyslane": 0, "Otwarcia": 0},
        ]}},
    ]
    out = build_md("Demo", demo_sections, DEFAULT_OUTPUT_DIR / "demo.md")
    print(f"Zapisano: {out}")
